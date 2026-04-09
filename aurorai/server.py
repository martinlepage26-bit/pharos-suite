from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form, Depends, Request, Response, Query
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from contextlib import asynccontextmanager
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import hashlib
import asyncio
import shutil
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import Any, Dict, List, Optional, Literal
import uuid
from datetime import datetime, timezone, timedelta
import fitz  # PyMuPDF
import aiofiles
import tempfile
import json
import re
from enum import Enum
from io import BytesIO
import socket
from urllib import request as urllib_request
from urllib.error import HTTPError, URLError
from zipfile import BadZipFile, ZipFile
from jose import JWTError, jwt
from pymongo import ReturnDocument

from compassai.backend.security import get_password_hash, verify_password

try:
    from docx import Document as WordDocument
except Exception:  # pragma: no cover - optional dependency
    WordDocument = None

try:
    from pdf2image import convert_from_path
except Exception:  # pragma: no cover - optional dependency
    convert_from_path = None

try:
    import pytesseract
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create uploads directory
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")
security = HTTPBearer(auto_error=False)
_session_secret = os.environ.get("AURORAI_SESSION_SECRET")
if not _session_secret:
    raise RuntimeError("AURORAI_SESSION_SECRET env var is required and must not be empty")
AURORAI_SESSION_SECRET = _session_secret
AURORAI_SESSION_COOKIE = os.environ.get("AURORAI_SESSION_COOKIE", "aurorai_session")
AURORAI_SESSION_TTL_MINUTES = max(int(os.environ.get("AURORAI_SESSION_TTL_MINUTES", "720") or "720"), 30)
AURORAI_SESSION_COOKIE_SECURE = os.environ.get("AURORAI_SESSION_COOKIE_SECURE", "false").strip().lower() in {"1", "true", "yes", "on"}
AURORAI_REQUIRE_OCR_RUNTIME = os.environ.get("AURORAI_REQUIRE_OCR_RUNTIME", "false").strip().lower() in {"1", "true", "yes", "on"}
AURORAI_JOB_POLL_INTERVAL_SECONDS = max(float(os.environ.get("AURORAI_JOB_POLL_INTERVAL_SECONDS", "0.5") or "0.5"), 0.1)
AURORAI_JOB_MAX_RETRIES = max(int(os.environ.get("AURORAI_JOB_MAX_RETRIES", "3") or "3"), 0)
AURORAI_WORKER_ID = os.environ.get("AURORAI_WORKER_ID", f"{socket.gethostname()}:{os.getpid()}")
_compassai_base_url = os.environ.get("COMPASSAI_BASE_URL")
if not _compassai_base_url:
    raise RuntimeError("COMPASSAI_BASE_URL env var is required — set to the CompassAI backend base URL (e.g. http://127.0.0.1:9205 locally)")
COMPASSAI_BASE_URL = _compassai_base_url.rstrip("/")
COMPASSAI_INGEST_TOKEN = os.environ.get("COMPASSAI_INGEST_TOKEN", "")
SUPPORTED_UPLOAD_EXTENSIONS = {".pdf", ".txt", ".docx"}
OCR_ENABLED = os.environ.get("AURORAI_ENABLE_OCR", "true").strip().lower() not in {"0", "false", "no", "off"}
PDF_TEXT_MIN_CHARS = max(int(os.environ.get("AURORAI_PDF_TEXT_MIN_CHARS", "40") or "40"), 1)
OCR_MAX_PAGES = max(int(os.environ.get("AURORAI_OCR_MAX_PAGES", "10") or "10"), 1)

# Categories
CATEGORIES = [
    "Academic Papers",
    "Invoices/Receipts",
    "Contracts",
    "Reports",
    "Personal Documents",
    "Resumes",
    "My Writings & Publications",
    "Uncategorized"
]

IDP_MISSION = (
    "AurorAI transforme des documents bloqués en données structurées, "
    "recherchables et exploitables, avec traçabilité, sécurité et conformité."
)

IDP_PIPELINE = [
    "Ingestion: upload / email / dépôt / API",
    "Lecture (OCR): extraction de texte depuis scans, images, PDF",
    "Compréhension (NLP): contexte, entités, relations",
    "Classification: type de document + sous-type",
    "Extraction: champs clés vers un schéma JSON/CSV avec score de confiance",
    "Contrôle (HITL): validation humaine ciblée sous seuil de confiance",
    "Routage & intégration: ERP/CRM/DMS, workflows et alertes",
    "Gouvernance: audit trail, versioning, rétention et contrôle d'accès",
]

class UserRole(str, Enum):
    OPERATOR = "operator"
    REVIEWER = "reviewer"
    ADMIN = "admin"


class ProcessingJobType(str, Enum):
    OCR = "ocr"
    CLASSIFY = "classify"
    SUMMARY = "summary"
    EXTRACT = "extract"
    CITATIONS = "citations"


class ProcessingJobStatus(str, Enum):
    QUEUED = "queued"
    RUNNING = "running"
    SUCCESS = "success"
    FAILED = "failed"


class UserCreate(BaseModel):
    email: str
    password: str
    name: str
    role: UserRole


class UserLogin(BaseModel):
    email: str
    password: str


class SessionToken(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: Dict[str, Any]


class InternalUser(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    role: UserRole
    hashed_password: str
    is_active: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class ProcessingJob(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    job_type: ProcessingJobType
    status: ProcessingJobStatus = ProcessingJobStatus.QUEUED
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    error_log: List[str] = Field(default_factory=list)
    retry_count: int = 0
    worker_id: Optional[str] = None
    requested_by_user_id: Optional[str] = None
    requested_by_role: Optional[str] = None
    output: Dict[str, Any] = Field(default_factory=dict)
    next_job_types: List[str] = Field(default_factory=list)


class ProcessingJobRequest(BaseModel):
    document_id: str
    job_type: ProcessingJobType
    next_job_types: List[ProcessingJobType] = Field(default_factory=list)


class ReviewDecisionRequest(BaseModel):
    decision: Literal["approve", "reject", "flag"]
    comment: str = ""


# Define Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    original_filename: str
    category: str = "Uncategorized"
    ai_suggested_category: Optional[str] = None
    document_type: Optional[str] = None
    classification_confidence: Optional[float] = None
    classification_rationale: Optional[str] = None
    file_size: int = 0
    page_count: int = 0
    text_preview: str = ""
    uploaded_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    is_academic: bool = False
    citations: List[str] = []
    summary: Optional[str] = None
    extraction: Dict[str, Any] = Field(default_factory=dict)
    control_checks: Dict[str, List[str]] = Field(default_factory=dict)
    review_required: bool = False
    compliance_note: Optional[str] = None
    source_hash: Optional[str] = None
    ingestion_details: Dict[str, Any] = Field(default_factory=dict)
    audit_log: List[Dict[str, Any]] = Field(default_factory=list)
    current_state: str = "uploaded"
    current_review_state: str = "pending"
    latest_run_id: Optional[str] = None
    latest_package_id: Optional[str] = None
    latest_handoff_id: Optional[str] = None
    latest_job_id: Optional[str] = None
    assigned_reviewer_user_id: Optional[str] = None
    processing_runs: List[Dict[str, Any]] = Field(default_factory=list)
    review_decision: Optional[Dict[str, Any]] = None
    review_decisions: List[Dict[str, Any]] = Field(default_factory=list)
    package_history: List[Dict[str, Any]] = Field(default_factory=list)
    handoff_history: List[Dict[str, Any]] = Field(default_factory=list)
    artifact_versions: List[Dict[str, Any]] = Field(default_factory=list)
    reading_list_id: Optional[str] = None
    tags: List[str] = []

class DocumentCreate(BaseModel):
    filename: str
    original_filename: str
    category: str = "Uncategorized"
    file_size: int = 0
    page_count: int = 0
    text_preview: str = ""

class DocumentUpdate(BaseModel):
    category: Optional[str] = None
    tags: Optional[List[str]] = None
    reading_list_id: Optional[str] = None


class DocumentExtractionError(RuntimeError):
    def __init__(self, message: str, *, status_code: int = 400):
        super().__init__(message)
        self.status_code = status_code
        self.message = message


def log_pipeline_event(event: str, **details: Any) -> None:
    payload = {
        "component": "aurorai",
        "event": event,
        **details,
    }
    logging.info(json.dumps(payload, sort_keys=True, ensure_ascii=True, default=str))


def build_runtime_capabilities() -> Dict[str, Any]:
    ocr_reason = []
    if not OCR_ENABLED:
        ocr_reason.append("disabled_by_feature_flag")
    if convert_from_path is None:
        ocr_reason.append("pdf2image_missing")
    if pytesseract is None:
        ocr_reason.append("pytesseract_missing")
    if shutil.which("pdftoppm") is None:
        ocr_reason.append("poppler_pdftoppm_missing")
    if shutil.which("tesseract") is None:
        ocr_reason.append("tesseract_binary_missing")

    ocr_available = OCR_ENABLED and not ocr_reason
    return {
        "supported_upload_extensions": sorted(SUPPORTED_UPLOAD_EXTENSIONS),
        "docx_enabled": WordDocument is not None,
        "ocr_enabled": OCR_ENABLED,
        "ocr_available": ocr_available,
        "ocr_reason": "available" if ocr_available else ", ".join(ocr_reason) or "disabled",
        "pdf_text_min_chars": PDF_TEXT_MIN_CHARS,
        "ocr_max_pages": OCR_MAX_PAGES,
    }


def text_quality(text: str, *, min_chars: int = PDF_TEXT_MIN_CHARS) -> str:
    stripped = (text or "").strip()
    if not stripped:
        return "none"

    compact = re.sub(r"\s+", "", stripped)
    alnum = re.sub(r"[^A-Za-z0-9]+", "", stripped)
    words = re.findall(r"[A-Za-z0-9]+", stripped)
    if len(alnum) < max(8, min_chars // 5):
        return "insufficient"
    if len(compact) < min_chars and len(words) < 3:
        return "insufficient"
    return "usable"


def extract_docx_text(docx_path: str) -> str:
    if WordDocument is None:
        raise DocumentExtractionError(
            "DOCX extraction is unavailable because python-docx is not installed.",
            status_code=503,
        )

    try:
        document = WordDocument(docx_path)
        blocks: List[str] = []
        blocks.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        blocks.append(cell_text)
        text = "\n".join(blocks).strip()
        if not text:
            raise DocumentExtractionError("DOCX extraction succeeded but no readable text was found.")
        log_pipeline_event(
            "docx_text_extraction_succeeded",
            path=docx_path,
            extracted_chars=len(text),
        )
        return text
    except DocumentExtractionError:
        raise
    except (BadZipFile, KeyError, ValueError) as exc:
        log_pipeline_event(
            "docx_text_extraction_failed",
            path=docx_path,
            error=str(exc),
        )
        raise DocumentExtractionError(f"DOCX extraction failed: {exc}") from exc
    except Exception as exc:
        log_pipeline_event(
            "docx_text_extraction_failed",
            path=docx_path,
            error=str(exc),
        )
        raise DocumentExtractionError(f"DOCX extraction failed: {exc}") from exc


def run_pdf_ocr(pdf_path: str, max_pages: int) -> Dict[str, Any]:
    capabilities = build_runtime_capabilities()
    if not capabilities["ocr_available"]:
        log_pipeline_event(
            "ocr_unavailable",
            path=pdf_path,
            reason=capabilities["ocr_reason"],
        )
        return {
            "text": "",
            "ocr_used": False,
            "ocr_status": "required_unavailable",
            "ocr_reason": capabilities["ocr_reason"],
            "warnings": [f"OCR required but unavailable: {capabilities['ocr_reason']}"],
        }

    try:
        images = convert_from_path(pdf_path, first_page=1, last_page=max_pages)
        ocr_text = []
        for index, image in enumerate(images, start=1):
            page_text = pytesseract.image_to_string(image)
            if page_text.strip():
                ocr_text.append(page_text)
            log_pipeline_event(
                "ocr_page_processed",
                path=pdf_path,
                page=index,
                extracted_chars=len(page_text.strip()),
            )

        text = "\n".join(ocr_text).strip()
        quality = text_quality(text)
        log_pipeline_event(
            "ocr_fallback_triggered",
            path=pdf_path,
            pages_processed=min(len(images), max_pages),
            extracted_chars=len(text),
            text_quality=quality,
        )
        return {
            "text": text,
            "ocr_used": True,
            "ocr_status": "performed" if text else "performed_no_text",
            "ocr_reason": "available",
            "warnings": [] if text else ["OCR completed but did not recover readable text."],
            "text_quality": quality,
        }
    except Exception as exc:
        log_pipeline_event(
            "ocr_failed",
            path=pdf_path,
            error=str(exc),
        )
        return {
            "text": "",
            "ocr_used": True,
            "ocr_status": "failed",
            "ocr_reason": str(exc),
            "warnings": [f"OCR fallback failed: {exc}"],
            "text_quality": "none",
        }


def extract_document_content(
    file_path: Path,
    *,
    raw_bytes: Optional[bytes] = None,
    max_pages: int = 10,
    allow_ocr: bool = True,
) -> Dict[str, Any]:
    file_ext = file_path.suffix.lower()

    if file_ext == ".txt":
        content = raw_bytes if raw_bytes is not None else file_path.read_bytes()
        text = content.decode("utf-8", errors="ignore")
        quality = text_quality(text, min_chars=1)
        log_pipeline_event(
            "text_file_extraction_used",
            path=str(file_path),
            extracted_chars=len(text.strip()),
        )
        return {
            "text": text,
            "page_count": 1,
            "text_source": "txt",
            "text_quality": quality,
            "ocr_status": "not_applicable",
            "ocr_reason": "not_applicable",
            "warnings": [],
        }

    if file_ext == ".docx":
        text = extract_docx_text(str(file_path))
        quality = text_quality(text, min_chars=1)
        return {
            "text": text,
            "page_count": 1,
            "text_source": "docx",
            "text_quality": quality,
            "ocr_status": "not_applicable",
            "ocr_reason": "not_applicable",
            "warnings": [],
        }

    if file_ext != ".pdf":
        raise DocumentExtractionError(f"File type {file_ext} is not supported for text extraction.")

    text, page_count = extract_pdf_text(str(file_path), max_pages=max_pages)
    quality = text_quality(text)
    if quality == "usable":
        log_pipeline_event(
            "pdf_text_extraction_succeeded",
            path=str(file_path),
            page_count=page_count,
            extracted_chars=len(text),
        )
        return {
            "text": text,
            "page_count": page_count,
            "text_source": "pdf_text",
            "text_quality": quality,
            "ocr_status": "not_needed",
            "ocr_reason": "fitz_text_sufficient",
            "warnings": [],
        }

    log_pipeline_event(
        "pdf_text_extraction_insufficient",
        path=str(file_path),
        page_count=page_count,
        extracted_chars=len(text),
        text_quality=quality,
    )
    if not allow_ocr:
        capabilities = build_runtime_capabilities()
        ocr_status = "queued" if capabilities["ocr_available"] else "required_unavailable"
        ocr_reason = "queued_for_async_ocr" if capabilities["ocr_available"] else capabilities["ocr_reason"]
        warnings = (
            ["OCR deferred to asynchronous processing."]
            if capabilities["ocr_available"]
            else [f"OCR required but unavailable: {capabilities['ocr_reason']}"]
        )
        return {
            "text": text.strip(),
            "page_count": page_count,
            "text_source": "pdf_text" if text.strip() else "pdf_scan_pending_ocr",
            "text_quality": quality,
            "ocr_status": ocr_status,
            "ocr_reason": ocr_reason,
            "warnings": warnings,
        }

    ocr_result = run_pdf_ocr(str(file_path), max_pages=min(page_count or max_pages, OCR_MAX_PAGES, max_pages))
    ocr_text = ocr_result.get("text", "").strip()
    ocr_quality = ocr_result.get("text_quality", text_quality(ocr_text))
    if ocr_quality == "usable":
        return {
            "text": ocr_text,
            "page_count": page_count,
            "text_source": "pdf_ocr",
            "text_quality": ocr_quality,
            "ocr_status": ocr_result["ocr_status"],
            "ocr_reason": ocr_result["ocr_reason"],
            "warnings": ocr_result.get("warnings", []),
        }

    return {
        "text": text.strip(),
        "page_count": page_count,
        "text_source": "pdf_text" if text.strip() else "pdf_unreadable",
        "text_quality": quality,
        "ocr_status": ocr_result["ocr_status"],
        "ocr_reason": ocr_result["ocr_reason"],
        "warnings": ocr_result.get("warnings", []),
    }


def get_document_text_for_processing(
    document: Dict[str, Any],
    *,
    max_pages: int = 10,
    allow_ocr: bool = False,
) -> Dict[str, Any]:
    stored_text = str(document.get("text_preview") or "").strip()
    ingestion_details = document.get("ingestion_details", {}) or {}
    if stored_text and ingestion_details.get("text_quality", "usable") == "usable":
        return {
            "text": stored_text,
            "page_count": document.get("page_count", 0),
            "text_source": ingestion_details.get("text_source", "stored_preview"),
            "text_quality": ingestion_details.get("text_quality", "usable"),
            "ocr_status": ingestion_details.get("ocr_status", "unknown"),
            "ocr_reason": ingestion_details.get("ocr_reason", "stored_preview"),
            "warnings": [],
        }

    file_path = UPLOAD_DIR / document["filename"]
    if not file_path.exists():
        return {
            "text": stored_text,
            "page_count": document.get("page_count", 0),
            "text_source": "missing_source_file",
            "text_quality": text_quality(stored_text),
            "ocr_status": ingestion_details.get("ocr_status", "unknown"),
            "ocr_reason": "source_file_missing",
            "warnings": ["Source file is missing; falling back to stored preview only."],
        }

    return extract_document_content(file_path, max_pages=max_pages, allow_ocr=allow_ocr)


def no_text_http_exception(result: Dict[str, Any], *, action: str) -> HTTPException:
    ocr_status = result.get("ocr_status")
    if ocr_status == "queued":
        return HTTPException(
            status_code=409,
            detail=f"OCR is queued before AurorA can {action} this document.",
        )
    if ocr_status == "required_unavailable":
        return HTTPException(
            status_code=503,
            detail=f"OCR is required to {action} this PDF, but the OCR runtime is unavailable ({result.get('ocr_reason')}).",
        )
    if ocr_status == "failed":
        return HTTPException(
            status_code=502,
            detail=f"OCR fallback failed while trying to {action} this PDF: {result.get('ocr_reason')}",
        )
    if ocr_status == "performed_no_text":
        return HTTPException(
            status_code=422,
            detail=f"OCR ran while trying to {action} this PDF, but no usable text was recovered.",
        )
    return HTTPException(status_code=400, detail=f"No text content available to {action}.")

class ReadingList(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: str = ""
    document_ids: List[str] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ReadingListCreate(BaseModel):
    name: str
    description: str = ""

class CategoryStats(BaseModel):
    category: str
    count: int
    percentage: float

class AICategorizationRequest(BaseModel):
    document_id: str

class AISummaryRequest(BaseModel):
    document_id: str

class CitationExtractionRequest(BaseModel):
    document_id: str


class EvidencePackageRequest(BaseModel):
    usecase_id: str
    producer: str = "aurorai"
    artifact_type: str = "evidence_package"


class EvidenceHandoffRequest(EvidencePackageRequest):
    compassai_base_url: Optional[str] = None


def normalize_email(email: str) -> str:
    return (email or "").strip().lower()


def build_role_capabilities(role: Optional[str]) -> Dict[str, bool]:
    normalized = (role or "").strip().lower()
    return {
        "can_upload": normalized in {UserRole.OPERATOR.value, UserRole.ADMIN.value},
        "can_process": normalized in {UserRole.OPERATOR.value, UserRole.ADMIN.value},
        "can_review": normalized in {UserRole.REVIEWER.value, UserRole.ADMIN.value},
        "can_admin": normalized == UserRole.ADMIN.value,
        "can_view_logs": normalized == UserRole.ADMIN.value,
        "can_batch_upload": normalized in {UserRole.OPERATOR.value, UserRole.ADMIN.value},
        "can_retry_jobs": normalized in {UserRole.REVIEWER.value, UserRole.ADMIN.value},
    }


def create_session_token(data: Dict[str, Any], expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=AURORAI_SESSION_TTL_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, AURORAI_SESSION_SECRET, algorithm="HS256")


def set_session_cookie(response: Response, token: str) -> None:
    response.set_cookie(
        key=AURORAI_SESSION_COOKIE,
        value=token,
        httponly=True,
        samesite="lax",
        secure=AURORAI_SESSION_COOKIE_SECURE,
        max_age=AURORAI_SESSION_TTL_MINUTES * 60,
        path="/",
    )


def clear_session_cookie(response: Response) -> None:
    response.delete_cookie(AURORAI_SESSION_COOKIE, path="/")


async def get_current_user(
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(security),
) -> Optional[Dict[str, Any]]:
    token = None
    if credentials and credentials.credentials:
        token = credentials.credentials
    elif request.cookies.get(AURORAI_SESSION_COOKIE):
        token = request.cookies.get(AURORAI_SESSION_COOKIE)

    if not token:
        return None

    try:
        payload = jwt.decode(token, AURORAI_SESSION_SECRET, algorithms=["HS256"])
        user_id = payload.get("sub")
        if not user_id:
            return None
        user = await db.users.find_one({"id": user_id, "is_active": True}, {"_id": 0, "hashed_password": 0})
        request.state.user = user
        return user
    except JWTError:
        return None


async def require_auth_user(
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(security),
) -> Dict[str, Any]:
    user = await get_current_user(request, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    request.state.user = user
    return user


def require_role(allowed_roles: List[UserRole]):
    async def _require_role(
        request: Request,
        credentials: HTTPAuthorizationCredentials = Depends(security),
    ) -> Dict[str, Any]:
        user = await require_auth_user(request, credentials)
        allowed = {role.value for role in allowed_roles}
        if user.get("role") not in allowed:
            raise HTTPException(status_code=403, detail="Not authorized")
        return user

    return _require_role


require_operator_access = require_role([UserRole.OPERATOR, UserRole.ADMIN])
require_reviewer_access = require_role([UserRole.REVIEWER, UserRole.ADMIN])
require_document_read_access = require_role([UserRole.OPERATOR, UserRole.REVIEWER, UserRole.ADMIN])
require_admin_access = require_role([UserRole.ADMIN])


def append_audit_event(document: Dict[str, Any], action: str, details: Dict[str, Any]) -> List[Dict[str, Any]]:
    events = document.get("audit_log", [])
    events.append(
        {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "action": action,
            "details": details,
        }
    )
    return events[-100:]


def latest_stage_run_id(document: Dict[str, Any], stage: str) -> Optional[str]:
    for run in reversed(document.get("processing_runs", [])):
        if run.get("stage") == stage:
            return run.get("id")
    return None


def append_processing_run(
    document: Dict[str, Any],
    stage: str,
    *,
    triggered_by: str,
    status: str,
    details: Optional[Dict[str, Any]] = None,
    job_id: Optional[str] = None,
    started_at: Optional[str] = None,
    completed_at: Optional[str] = None,
) -> tuple[List[Dict[str, Any]], Dict[str, Any]]:
    runs = list(document.get("processing_runs", []))
    stage_runs = [run for run in runs if run.get("stage") == stage]
    run_doc = {
        "id": str(uuid.uuid4()),
        "job_id": job_id,
        "job_type": stage,
        "stage": stage,
        "status": status,
        "triggered_by": triggered_by,
        "parent_run_id": latest_stage_run_id(document, stage),
        "iteration_index": len(stage_runs) + 1,
        "started_at": started_at or iso_now(),
        "completed_at": completed_at or iso_now(),
        "timestamp": completed_at or iso_now(),
        "details": details or {},
        "output": details or {},
    }
    runs.append(run_doc)
    return runs[-150:], run_doc


def append_review_decision(
    document: Dict[str, Any],
    *,
    run_id: Optional[str],
    decision_type: str,
    rationale: str,
    resulting_state: str,
    actor_type: str = "system",
    actor_id: str = "aurorai",
    reviewer_user_id: Optional[str] = None,
    comment: Optional[str] = None,
) -> tuple[List[Dict[str, Any]], Dict[str, Any]]:
    decisions = list(document.get("review_decisions", []))
    decision_doc = {
        "id": str(uuid.uuid4()),
        "run_id": run_id,
        "parent_decision_id": decisions[-1]["id"] if decisions else None,
        "review_round": len(decisions) + 1,
        "actor_type": actor_type,
        "actor_id": actor_id,
        "decision_type": decision_type,
        "decision": decision_type,
        "reviewer_user_id": reviewer_user_id,
        "rationale": rationale,
        "comment": comment or rationale,
        "resulting_state": resulting_state,
        "created_at": iso_now(),
        "timestamp": iso_now(),
    }
    decisions.append(decision_doc)
    return decisions[-150:], decision_doc


def append_package_history(
    document: Dict[str, Any],
    *,
    evidence_id: str,
    version: int,
    usecase_id: str,
    package_hash: str,
) -> tuple[List[Dict[str, Any]], Dict[str, Any]]:
    history = list(document.get("package_history", []))
    package_doc = {
        "id": evidence_id,
        "version": version,
        "usecase_id": usecase_id,
        "hash": package_hash,
        "created_at": iso_now(),
        "supersedes_package_id": history[-1]["id"] if history else None,
    }
    history.append(package_doc)
    return history[-100:], package_doc


def append_handoff_history(
    document: Dict[str, Any],
    *,
    evidence_id: str,
    target: str,
    target_url: str,
    status_code: int,
) -> tuple[List[Dict[str, Any]], Dict[str, Any]]:
    history = list(document.get("handoff_history", []))
    handoff_doc = {
        "id": str(uuid.uuid4()),
        "evidence_id": evidence_id,
        "handoff_id": str(uuid.uuid4()),
        "target": target,
        "target_system": target,
        "target_url": target_url,
        "status_code": status_code,
        "status": "succeeded" if status_code < 400 else "failed",
        "created_at": iso_now(),
        "timestamp": iso_now(),
        "payload": {
            "evidence_id": evidence_id,
            "target_url": target_url,
            "status_code": status_code,
        },
    }
    history.append(handoff_doc)
    return history[-100:], handoff_doc


def append_artifact_version(
    document: Dict[str, Any],
    *,
    source: str,
    data: Dict[str, Any],
) -> tuple[List[Dict[str, Any]], Dict[str, Any]]:
    versions = list(document.get("artifact_versions", []))
    version_doc = {
        "version": len(versions) + 1,
        "created_at": iso_now(),
        "source": source,
        "data": data,
    }
    versions.append(version_doc)
    return versions[-150:], version_doc


def parse_json_response(raw: str) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
        if not match:
            return None
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            return None


def run_control_checks(extraction_fields: Dict[str, Dict[str, Any]]) -> Dict[str, List[str]]:
    missing = []
    anomalies = []
    invalid = []
    duplicates = []

    values_seen = {}
    for field_name, field_data in extraction_fields.items():
        value = str(field_data.get("value", "")).strip()
        if not value:
            missing.append(field_name)
            continue

        if value in values_seen:
            duplicates.append(f"{field_name} duplicates {values_seen[value]}")
        else:
            values_seen[value] = field_name

        lower_field = field_name.lower()
        if "date" in lower_field and not re.search(r"\d{4}-\d{2}-\d{2}|\d{2}/\d{2}/\d{4}", value):
            invalid.append(f"{field_name} has non-standard date format: {value}")
        if any(token in lower_field for token in ["amount", "total", "price"]) and not re.search(r"\d", value):
            invalid.append(f"{field_name} has no numeric value: {value}")
        if "id" in lower_field or "number" in lower_field:
            if len(value) < 4:
                anomalies.append(f"{field_name} looks too short: {value}")

    return {
        "missing_fields": missing,
        "inconsistencies": anomalies,
        "duplicate_values": duplicates,
        "invalid_values": invalid,
    }


def compute_sha256_bytes(data: bytes) -> str:
    return f"sha256:{hashlib.sha256(data).hexdigest()}"


def compute_sha256_json(data: Any) -> str:
    canonical = json.dumps(data, sort_keys=True, separators=(",", ":"), ensure_ascii=True).encode("utf-8")
    return compute_sha256_bytes(canonical)


def detect_pii_markers(text: str) -> List[str]:
    markers = []
    pii_patterns = {
        "ssn": r"\b\d{3}-\d{2}-\d{4}\b",
        "credit_card": r"\b(?:\d[ -]?){13,16}\b",
        "email": r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b",
        "phone": r"\b(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)\d{3}[-.\s]?\d{4}\b",
        "iban": r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b",
    }

    for label, pattern in pii_patterns.items():
        if re.search(pattern, text, flags=re.IGNORECASE):
            markers.append(label)

    lowered = text.lower()
    keyword_markers = {
        "financial_context": ["account number", "routing number", "wire transfer"],
        "medical_context": ["patient", "diagnosis", "medical record", "health card"],
    }
    for label, keywords in keyword_markers.items():
        if any(keyword in lowered for keyword in keywords):
            markers.append(label)

    return sorted(set(markers))


def get_below_threshold_fields(
    extraction_fields: Dict[str, Dict[str, Any]],
    threshold: float = 0.85,
) -> List[Dict[str, Any]]:
    below_threshold = []
    for field_name, field_data in extraction_fields.items():
        confidence = field_data.get("confidence")
        if isinstance(confidence, (int, float)) and float(confidence) < threshold:
            below_threshold.append(
                {
                    "field": field_name,
                    "confidence": round(float(confidence), 4),
                    "threshold": threshold,
                }
            )
    return below_threshold


def detect_high_value_transaction(extraction_fields: Dict[str, Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    for field_name, field_data in extraction_fields.items():
        if not any(token in field_name.lower() for token in ["amount", "total", "price", "payment"]):
            continue

        raw_value = str(field_data.get("value", "")).strip()
        normalized = re.sub(r"[^0-9.,-]", "", raw_value).replace(",", "")
        if not normalized:
            continue

        try:
            amount = float(normalized)
        except ValueError:
            continue

        if amount >= 10000:
            return {
                "field": field_name,
                "amount": amount,
            }

    return None


def build_hitl_reason(
    below_threshold_fields: List[Dict[str, Any]],
    pii_markers: List[str],
    high_value_trigger: Optional[Dict[str, Any]],
) -> Optional[str]:
    reasons = []
    if below_threshold_fields:
        reasons.append("confidence_below_threshold")
    if pii_markers:
        reasons.append("pii_detected")
    if high_value_trigger:
        reasons.append("high_value_transaction")
    return ", ".join(reasons) if reasons else None


def iso_now() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def build_evidence_package(document: Dict[str, Any], request: EvidencePackageRequest) -> Dict[str, Any]:
    extraction = document.get("extraction", {}) or {}
    fields = extraction.get("fields", {}) if isinstance(extraction.get("fields"), dict) else {}
    controls = document.get("control_checks", {}) or {}

    below_threshold_fields = get_below_threshold_fields(fields)
    pii_markers = detect_pii_markers(document.get("text_preview", ""))
    high_value_trigger = detect_high_value_transaction(fields)
    mandatory_fields_present = not bool(controls.get("missing_fields"))
    hitl_trigger_reason = build_hitl_reason(below_threshold_fields, pii_markers, high_value_trigger)
    hitl_triggered = bool(document.get("review_required") or hitl_trigger_reason)
    validation_status = (
        "awaiting_hitl"
        if hitl_triggered
        else "auto_approved"
        if mandatory_fields_present and not controls.get("invalid_values")
        else "needs_review"
    )

    payload = {
        "evidence_package": {
            "extraction_id": f"EXT-{datetime.now(timezone.utc).year}-{document['id'][:8].upper()}",
            "schema_version": "2026-03-01",
            "document_metadata": {
                "source_hash": document.get("source_hash"),
                "document_type": extraction.get("document_type") or document.get("document_type") or "unknown",
                "processing_timestamp": iso_now(),
                "original_filename": document.get("original_filename"),
            },
            "extraction_results": {
                "fields": fields,
                "mandatory_fields_present": mandatory_fields_present,
                "below_threshold_fields": below_threshold_fields,
            },
            "quality_controls": {
                "pii_detected": pii_markers,
                "pii_masking_applied": False,
                "hitl_triggered": hitl_triggered,
                "hitl_trigger_reason": hitl_trigger_reason,
                "confidence_check": "passed" if not below_threshold_fields else "failed",
                "validation_status": validation_status,
            },
            "metrics": {
                "processing_time_ms": None,
                "model_version": os.environ.get("AURORAI_MODEL_VERSION", "aurorai-preview"),
                "accuracy_score_benchmark": None,
                "benchmark_dataset": None,
                "benchmark_date": None,
            },
            "audit_trail": {
                "processor_id": os.environ.get("AURORAI_PROCESSOR_ID", "aurorai-api"),
                "validation_chain": [event.get("action") for event in document.get("audit_log", []) if event.get("action")],
                "approver": None,
                "override_applied": False,
            },
            "if_trace_receipt": None,
            "if_trace_binding_note": (
                "if.trace binding is available via operator-initiated receipt generation; "
                "not automatic on every package in the current implementation"
            ),
        }
    }

    envelope_control_checks = {
        "missing_fields": controls.get("missing_fields", []),
        "invalid_values": controls.get("invalid_values", []),
        "duplicate_values": controls.get("duplicate_values", []),
        "inconsistencies": controls.get("inconsistencies", []),
        "review_required": hitl_triggered,
        "pii_detected": pii_markers,
        "high_value_trigger": high_value_trigger,
    }

    return {
        "usecase_id": request.usecase_id,
        "producer": request.producer,
        "artifact_type": request.artifact_type,
        "payload": payload,
        "hash": compute_sha256_json(payload),
        "control_checks": envelope_control_checks,
    }


def heuristic_document_type(text: str) -> str:
    lowered = text.lower()
    if "invoice" in lowered or "total amount" in lowered:
        return "invoice"
    if "agreement" in lowered or "contract" in lowered:
        return "contract"
    if "resume" in lowered or "curriculum vitae" in lowered:
        return "resume"
    if "report" in lowered:
        return "report"
    return "text_document"


def heuristic_extract_fields(text: str) -> Dict[str, Any]:
    field_patterns = [
        ("invoice_number", r"(?:invoice(?: number| no\.?| #)?)[\s:.-]*([A-Z0-9-]+)", 0.72),
        ("total_amount", r"(?:total(?: amount)?|amount due)[\s:.-]*([$]?\d[\d,]*(?:\.\d{2})?)", 0.72),
        ("vendor_name", r"(?:vendor|supplier|from)[\s:.-]*([^\n]+)", 0.68),
        ("document_date", r"(?:date)[\s:.-]*([0-9]{4}-[0-9]{2}-[0-9]{2}|[0-9]{2}/[0-9]{2}/[0-9]{4}|[A-Z][a-z]{2,8}\s+\d{1,2},\s+\d{4})", 0.66),
        ("contact_email", r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", 0.7),
        ("contact_phone", r"((?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)\d{3}[-.\s]?\d{4})", 0.67),
    ]

    fields: Dict[str, Dict[str, Any]] = {}
    for field_name, pattern, confidence in field_patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if not match:
            continue
        value = match.group(1).strip()
        evidence_snippet = match.group(0).strip()[:120]
        fields[field_name] = {
            "value": value,
            "confidence": confidence,
            "evidence": evidence_snippet,
        }

    first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
    if first_line and "document_title" not in fields:
        fields["document_title"] = {
            "value": first_line[:120],
            "confidence": 0.58,
            "evidence": first_line[:120],
        }

    return {
        "document_type": heuristic_document_type(text),
        "fields": fields,
        "missing_or_ambiguous": [] if fields else ["No structured fields were identified heuristically"],
        "recommended_next_actions": (
            ["Run HITL validation", "Configure PHAROS_LLM_KEY for richer extraction"]
            if fields
            else ["Run HITL validation", "Provide a more structured source document"]
        ),
    }


async def post_json(url: str, payload: Dict[str, Any], headers: Dict[str, str]) -> Dict[str, Any]:
    data = json.dumps(payload).encode("utf-8")

    def _send() -> Dict[str, Any]:
        req = urllib_request.Request(
            url,
            data=data,
            headers={"Content-Type": "application/json", **headers},
            method="POST",
        )
        try:
            with urllib_request.urlopen(req, timeout=20) as response:
                raw_body = response.read().decode("utf-8")
                return {
                    "status_code": response.getcode(),
                    "body": json.loads(raw_body) if raw_body else None,
                }
        except HTTPError as exc:
            raw_body = exc.read().decode("utf-8")
            return {
                "status_code": exc.code,
                "body": json.loads(raw_body) if raw_body else None,
            }
        except URLError as exc:
            raise HTTPException(status_code=502, detail=f"CompassAI handoff failed: {exc.reason}") from exc

    return await asyncio.to_thread(_send)


async def create_evidence_record(doc_id: str, request: EvidencePackageRequest) -> Dict[str, Any]:
    doc = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    extraction = doc.get("extraction", {}) or {}
    fields = extraction.get("fields", {}) if isinstance(extraction.get("fields"), dict) else {}
    if not fields:
        raise HTTPException(
            status_code=400,
            detail="Structured extraction is required before an evidence package can be generated.",
        )

    if not doc.get("source_hash"):
        file_path = UPLOAD_DIR / doc["filename"]
        if file_path.exists():
            with open(file_path, "rb") as handle:
                source_hash = compute_sha256_bytes(handle.read())
            doc["source_hash"] = source_hash
            await db.documents.update_one({"id": doc_id}, {"$set": {"source_hash": source_hash}})
        else:
            raise HTTPException(status_code=400, detail="Source hash is missing and the source file is no longer available.")

    evidence_envelope = build_evidence_package(doc, request)
    version = await db.evidence_packages.count_documents(
        {"document_id": doc_id, "usecase_id": request.usecase_id}
    ) + 1
    evidence_id = str(uuid.uuid4())
    created_at = iso_now()

    evidence_record = {
        "id": evidence_id,
        "document_id": doc_id,
        "version": version,
        "created_at": created_at,
        **evidence_envelope,
    }

    await db.evidence_packages.insert_one(dict(evidence_record))
    package_history, package_doc = append_package_history(
        doc,
        evidence_id=evidence_id,
        version=version,
        usecase_id=request.usecase_id,
        package_hash=evidence_record["hash"],
    )
    updated_audit_log = append_audit_event(
        doc,
        "evidence_package",
        {
            "evidence_id": evidence_id,
            "usecase_id": request.usecase_id,
            "version": version,
        },
    )
    await db.documents.update_one(
        {"id": doc_id},
        {
            "$set": {
                "audit_log": updated_audit_log,
                "package_history": package_history,
                "latest_package_id": package_doc["id"],
                "current_state": "packaged",
            }
        },
    )
    evidence_record["document_audit_log"] = updated_audit_log
    return evidence_record

# Helper function to extract text from PDF
def extract_pdf_text(pdf_path: str, max_pages: int = 10) -> tuple[str, int]:
    """Extract text from PDF file"""
    try:
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        text = ""
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            text += page.get_text()
        doc.close()
        return text.strip(), page_count
    except Exception as e:
        log_pipeline_event(
            "pdf_text_extraction_failed",
            path=pdf_path,
            error=str(e),
        )
        logging.error(f"Error extracting PDF text: {e}")
        return "", 0

# Helper function to extract citations
def extract_citations_from_text(text: str) -> List[str]:
    """Extract citations from text using common patterns"""
    citations = []
    # Pattern for common citation formats
    patterns = [
        r'\[(\d+)\]',  # [1], [2], etc.
        r'\(([A-Z][a-z]+(?:\s+(?:et\s+al\.?|&\s+[A-Z][a-z]+)?)?,?\s*\d{4})\)',  # (Author, 2020)
        r'([A-Z][a-z]+(?:\s+(?:et\s+al\.?|&\s+[A-Z][a-z]+)?)\s*\(\d{4}\))',  # Author (2020)
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        citations.extend(matches[:20])  # Limit to 20 citations per pattern
    
    return list(set(citations))[:50]  # Return unique citations, max 50

# AI Integration using Pharos
def get_llm_api_key() -> str:
    return os.environ.get("OPENAI_API_KEY") or os.environ.get("PHAROS_LLM_KEY") or ""


def get_llm_base_url() -> Optional[str]:
    return os.environ.get("OPENAI_BASE_URL") or None


def infer_category_from_document_type(document_type: str) -> str:
    mapping = {
        "invoice": "Invoices/Receipts",
        "contract": "Contracts",
        "resume": "Resumes",
        "report": "Reports",
        "text_document": "Uncategorized",
        "unknown": "Uncategorized",
    }
    return mapping.get((document_type or "unknown").lower(), "Uncategorized")


async def run_llm_prompt(system_message: str, prompt: str, model: str = "gpt-4o-mini") -> str:
    from pharos_integrations.llm.chat import LlmChat, SystemMessage, UserMessage

    api_key = get_llm_api_key()
    if not api_key:
        raise RuntimeError("No OpenAI-compatible LLM key configured.")

    chat = LlmChat(
        model=model,
        api_key=api_key,
        base_url=get_llm_base_url(),
    )
    return await asyncio.to_thread(
        chat.chat,
        [
            SystemMessage(content=system_message),
            UserMessage(content=prompt),
        ],
    )


async def ai_categorize_document(text: str) -> Dict[str, Any]:
    """Use AI to categorize document based on content."""
    try:
        if not get_llm_api_key():
            document_type = heuristic_document_type(text)
            log_pipeline_event(
                "classification_heuristic_used",
                reason="llm_key_missing",
                document_type=document_type,
            )
            return {
                "category": infer_category_from_document_type(document_type),
                "document_type": document_type,
                "confidence": 0.42 if document_type != "unknown" else 0.0,
                "rationale": "No LLM key configured. Heuristic categorization was used.",
            }

        response = await run_llm_prompt(
            """You are AurorAI, an Intelligent Document Processing (IDP) engine.
Goal: turn documents into structured, actionable data with compliance-minded care.

Task: classify the document into ONE category from the allowed list AND detect a more specific document_type.

Return ONLY valid JSON with keys:
- category (must match exactly one allowed category)
- document_type (short string, e.g., "invoice", "employment_contract", "medical_intake_form", "loan_application", "court_filing")
- confidence (0.0–1.0)
- rationale (max 2 sentences)

Allowed categories:
Academic Papers; Invoices/Receipts; Contracts; Reports; Personal Documents; Resumes; My Writings & Publications; Uncategorized""",
            f"Classify this document:\n\n{text[:3000] if len(text) > 3000 else text}",
        )
        log_pipeline_event(
            "classification_llm_used",
            model="gpt-4o-mini",
            prompt_chars=min(len(text), 3000),
        )

        payload = parse_json_response(response.strip()) or {}
        category = payload.get("category", "Uncategorized")
        document_type = payload.get("document_type", "unknown")
        confidence = float(payload.get("confidence", 0.0) or 0.0)
        confidence = max(0.0, min(1.0, confidence))
        rationale = str(payload.get("rationale", "Model response could not be validated."))

        if category in CATEGORIES:
            return {
                "category": category,
                "document_type": document_type,
                "confidence": confidence,
                "rationale": rationale,
            }

        for cat in CATEGORIES:
            if cat.lower() in str(response).lower():
                return {
                    "category": cat,
                    "document_type": document_type,
                    "confidence": confidence,
                    "rationale": rationale,
                }

        return {
            "category": infer_category_from_document_type(document_type),
            "document_type": document_type,
            "confidence": confidence,
            "rationale": rationale,
        }
    except Exception as e:
        logging.error(f"AI categorization error: {e}")
        document_type = heuristic_document_type(text)
        log_pipeline_event(
            "classification_heuristic_used",
            reason=f"llm_error:{str(e)}",
            document_type=document_type,
        )
        return {
            "category": infer_category_from_document_type(document_type),
            "document_type": document_type,
            "confidence": 0.35 if document_type != "unknown" else 0.0,
            "rationale": f"LLM categorization unavailable. Heuristic fallback used: {str(e)}",
        }


async def ai_generate_summary(text: str) -> str:
    """Use AI to generate document summary"""
    try:
        if not get_llm_api_key():
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            preview = lines[:4]
            if not preview:
                return "Summary not available."
            return "Purpose\n- Review uploaded document\n\nKey points\n- " + "\n- ".join(preview[:3])

        response = await run_llm_prompt(
            """You are AurorAI, an IDP assistant. Summarize for operational use.

Output format:
- Purpose (1 sentence)
- Key points (3–6 bullets)
- Decisions / obligations / deadlines (bullets, if any)
- Risks or missing info (bullets, if any)

Keep it under 180 words. No fluff.""",
            f"Summarize this document:\n\n{text[:5000] if len(text) > 5000 else text}",
        )
        return response.strip()
    except Exception as e:
        logging.error(f"AI summary error: {e}")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        preview = lines[:4]
        if not preview:
            return "Summary generation failed."
        return "Purpose\n- Review uploaded document\n\nKey points\n- " + "\n- ".join(preview[:3])


async def ai_extract_fields(text: str) -> Dict[str, Any]:
    """Use AI to extract structured fields with confidence scores."""
    try:
        if not get_llm_api_key():
            extracted = heuristic_extract_fields(text)
            extracted["missing_or_ambiguous"] = extracted.get("missing_or_ambiguous", []) + ["LLM extraction unavailable - heuristic fallback used"]
            return extracted

        response = await run_llm_prompt(
            """You are AurorAI, an Intelligent Document Processing (IDP) extractor.

Extract key fields as structured data. Return ONLY valid JSON:
{
  "document_type": "...",
  "fields": {
     "<field_name>": {"value": "...", "confidence": 0.0-1.0, "evidence": "short quote"}
  },
  "missing_or_ambiguous": ["..."],
  "recommended_next_actions": ["..."]
}

Rules:
- Prefer accuracy over completeness.
- If a field is not present, do NOT guess; add it to missing_or_ambiguous.
- Evidence must be a short snippet from the text (max ~12 words).""",
            f"Extract fields from:\n\n{text[:6000] if len(text) > 6000 else text}",
        )
        parsed = parse_json_response(response.strip())
        if not parsed or not parsed.get("fields"):
            extracted = heuristic_extract_fields(text)
            extracted["missing_or_ambiguous"] = extracted.get("missing_or_ambiguous", []) + ["LLM response was empty or invalid - heuristic fallback used"]
            return extracted
        return parsed
    except Exception as e:
        logging.error(f"AI extraction error: {e}")
        extracted = heuristic_extract_fields(text)
        extracted["missing_or_ambiguous"] = extracted.get("missing_or_ambiguous", []) + [f"Extraction failed - heuristic fallback used: {str(e)}"]
        return extracted


async def ai_extract_citations(text: str) -> List[str]:
    """Use AI to extract and format citations"""
    try:
        if not get_llm_api_key():
            return extract_citations_from_text(text)

        response = await run_llm_prompt(
            """You are a citation extraction expert. Extract all references and citations from the document text.
Format each citation on a new line. Include author names, year, and title when available.
Return ONLY the citations, one per line, no numbering or bullets.""",
            f"Extract all citations from this document:\n\n{text[:6000] if len(text) > 6000 else text}",
        )
        citations = [c.strip() for c in response.strip().split('\n') if c.strip()]
        return citations[:50]
    except Exception as e:
        logging.error(f"AI citation extraction error: {e}")
        return extract_citations_from_text(text)


def serialize_user(user: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not user:
        return None
    role = user.get("role")
    return {
        "id": user.get("id"),
        "email": user.get("email"),
        "name": user.get("name"),
        "role": role,
        "capabilities": build_role_capabilities(role),
    }


async def ensure_bootstrap_user(email: str, password: str, name: str, role: UserRole) -> None:
    normalized_email = normalize_email(email)
    if not normalized_email or not password:
        return

    existing = await db.users.find_one({"email": normalized_email}, {"_id": 0})
    if existing:
        return

    user = InternalUser(
        email=normalized_email,
        name=name,
        role=role,
        hashed_password=get_password_hash(password),
    )
    user_doc = user.model_dump()
    user_doc["created_at"] = user_doc["created_at"].isoformat()
    await db.users.insert_one(user_doc)


async def seed_bootstrap_users() -> None:
    await ensure_bootstrap_user(
        os.environ.get("AURORAI_BOOTSTRAP_ADMIN_EMAIL", ""),
        os.environ.get("AURORAI_BOOTSTRAP_ADMIN_PASSWORD", ""),
        os.environ.get("AURORAI_BOOTSTRAP_ADMIN_NAME", "AurorA Admin"),
        UserRole.ADMIN,
    )
    await ensure_bootstrap_user(
        os.environ.get("AURORAI_BOOTSTRAP_OPERATOR_EMAIL", ""),
        os.environ.get("AURORAI_BOOTSTRAP_OPERATOR_PASSWORD", ""),
        os.environ.get("AURORAI_BOOTSTRAP_OPERATOR_NAME", "AurorA Operator"),
        UserRole.OPERATOR,
    )
    await ensure_bootstrap_user(
        os.environ.get("AURORAI_BOOTSTRAP_REVIEWER_EMAIL", ""),
        os.environ.get("AURORAI_BOOTSTRAP_REVIEWER_PASSWORD", ""),
        os.environ.get("AURORAI_BOOTSTRAP_REVIEWER_NAME", "AurorA Reviewer"),
        UserRole.REVIEWER,
    )


async def backfill_document_schema() -> None:
    documents = await db.documents.find({}, {"_id": 0}).to_list(5000)
    for document in documents:
        updates: Dict[str, Any] = {}
        defaults = {
            "processing_runs": [],
            "artifact_versions": [],
            "review_decision": None,
            "review_decisions": [],
            "handoff_history": [],
            "package_history": [],
            "current_state": "uploaded",
            "current_review_state": "pending",
            "assigned_reviewer_user_id": None,
            "latest_job_id": None,
        }
        for key, default_value in defaults.items():
            if key not in document:
                updates[key] = default_value
        if updates:
            await db.documents.update_one({"id": document["id"]}, {"$set": updates})


def enforce_required_runtime() -> None:
    capabilities = build_runtime_capabilities()
    if AURORAI_REQUIRE_OCR_RUNTIME and not capabilities["ocr_available"]:
        raise RuntimeError(
            f"AurorAI requires OCR runtime in this environment, but it is unavailable ({capabilities['ocr_reason']})."
        )


async def persist_processing_run_record(document_id: str, run_doc: Dict[str, Any]) -> None:
    payload = {"document_id": document_id, **run_doc}
    await db.processing_runs.insert_one(dict(payload))


async def append_document_audit(document_id: str, action: str, details: Dict[str, Any]) -> None:
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        return
    await db.documents.update_one(
        {"id": document_id},
        {"$set": {"audit_log": append_audit_event(document, action, details)}},
    )


async def enqueue_processing_job(
    document_id: str,
    job_type: ProcessingJobType,
    user: Optional[Dict[str, Any]],
    *,
    next_job_types: Optional[List[ProcessingJobType]] = None,
) -> Dict[str, Any]:
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    existing_jobs = await db.processing_jobs.find(
        {"document_id": document_id, "job_type": job_type.value},
        {"_id": 0},
    ).to_list(500)
    job = ProcessingJob(
        document_id=document_id,
        job_type=job_type,
        retry_count=len(existing_jobs),
        requested_by_user_id=user.get("id") if user else None,
        requested_by_role=user.get("role") if user else None,
        next_job_types=[item.value for item in (next_job_types or [])],
    )
    job_doc = job.model_dump()
    if isinstance(job_doc.get("job_type"), Enum):
        job_doc["job_type"] = job_doc["job_type"].value
    if isinstance(job_doc.get("status"), Enum):
        job_doc["status"] = job_doc["status"].value
    await db.processing_jobs.insert_one(dict(job_doc))

    await db.documents.update_one(
        {"id": document_id},
        {
            "$set": {
                "latest_job_id": job_doc["id"],
                "current_state": "processing_queued",
                "audit_log": append_audit_event(
                    document,
                    "job_enqueued",
                    {
                        "job_id": job_doc["id"],
                        "job_type": job_doc["job_type"],
                        "requested_by_user_id": job_doc.get("requested_by_user_id"),
                    },
                ),
            }
        },
    )
    return job_doc


async def claim_next_processing_job() -> Optional[Dict[str, Any]]:
    return await db.processing_jobs.find_one_and_update(
        {"status": ProcessingJobStatus.QUEUED.value},
        {
            "$set": {
                "status": ProcessingJobStatus.RUNNING.value,
                "started_at": iso_now(),
                "worker_id": AURORAI_WORKER_ID,
            }
        },
        sort=[("created_at", 1)],
        return_document=ReturnDocument.AFTER,
    )


async def mark_job_result(
    job: Dict[str, Any],
    *,
    status: ProcessingJobStatus,
    output: Optional[Dict[str, Any]] = None,
    error_message: Optional[str] = None,
) -> None:
    update_payload: Dict[str, Any] = {
        "status": status.value,
        "completed_at": iso_now(),
        "output": output or {},
    }
    if error_message:
        update_payload["error_log"] = list(job.get("error_log", [])) + [error_message]

    await db.processing_jobs.update_one({"id": job["id"]}, {"$set": update_payload})
    if status == ProcessingJobStatus.FAILED:
        document = await db.documents.find_one({"id": job["document_id"]}, {"_id": 0})
        if document:
            await db.documents.update_one(
                {"id": job["document_id"]},
                {
                    "$set": {
                        "current_state": "processing_failed",
                        "latest_job_id": job["id"],
                        "audit_log": append_audit_event(
                            document,
                            "job_failed",
                            {
                                "job_id": job["id"],
                                "job_type": job["job_type"],
                                "error": error_message,
                            },
                        ),
                    }
                },
            )


async def process_ocr_job(document: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    file_path = UPLOAD_DIR / document["filename"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on server")

    extraction_result = extract_document_content(file_path, max_pages=OCR_MAX_PAGES, allow_ocr=True)
    text = extraction_result.get("text", "").strip()
    if extraction_result.get("text_quality") != "usable":
        raise no_text_http_exception(extraction_result, action="recover text from")

    processing_runs, run_doc = append_processing_run(
        document,
        ProcessingJobType.OCR.value,
        triggered_by="job_worker",
        status="completed",
        job_id=job["id"],
        started_at=job.get("started_at"),
        completed_at=iso_now(),
        details={
            "text_source": extraction_result.get("text_source"),
            "ocr_status": extraction_result.get("ocr_status"),
            "extracted_chars": len(text),
        },
    )
    artifact_versions, _ = append_artifact_version(
        document,
        source="ocr",
        data={
            "text_preview": text[:2000],
            "ingestion_details": {
                "text_source": extraction_result.get("text_source"),
                "text_quality": extraction_result.get("text_quality"),
                "ocr_status": extraction_result.get("ocr_status"),
                "ocr_reason": extraction_result.get("ocr_reason"),
                "warnings": extraction_result.get("warnings", []),
            },
        },
    )
    updated_document = {
        "text_preview": text[:2000],
        "page_count": extraction_result.get("page_count", document.get("page_count", 0)),
        "ingestion_details": {
            "text_source": extraction_result.get("text_source"),
            "text_quality": extraction_result.get("text_quality"),
            "ocr_status": extraction_result.get("ocr_status"),
            "ocr_reason": extraction_result.get("ocr_reason"),
            "warnings": extraction_result.get("warnings", []),
        },
        "current_state": "ocr_complete",
        "current_review_state": "pending",
        "review_required": False,
        "processing_runs": processing_runs,
        "artifact_versions": artifact_versions,
        "latest_run_id": run_doc["id"],
        "latest_job_id": job["id"],
        "audit_log": append_audit_event(
            document,
            "ocr",
            {
                "job_id": job["id"],
                "ocr_status": extraction_result.get("ocr_status"),
                "text_source": extraction_result.get("text_source"),
            },
        ),
    }
    await db.documents.update_one({"id": document["id"]}, {"$set": updated_document})
    await persist_processing_run_record(document["id"], run_doc)
    return {"document_id": document["id"], "ocr_status": extraction_result.get("ocr_status")}


async def process_classify_job(document: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    text_result = get_document_text_for_processing(document, max_pages=10)
    text = text_result.get("text", "")
    if not text or text_result.get("text_quality") != "usable":
        raise no_text_http_exception(text_result, action="categorize")

    classification = await ai_categorize_document(text)
    suggested_category = classification["category"]
    review_required = classification.get("confidence", 0.0) < 0.75
    processing_runs, run_doc = append_processing_run(
        document,
        ProcessingJobType.CLASSIFY.value,
        triggered_by="job_worker",
        status="completed",
        job_id=job["id"],
        started_at=job.get("started_at"),
        completed_at=iso_now(),
        details={
            "category": suggested_category,
            "document_type": classification.get("document_type"),
            "confidence": classification.get("confidence"),
        },
    )
    review_decisions = list(document.get("review_decisions", []))
    latest_review_decision = document.get("review_decision")
    current_review_state = "classification_complete"
    if review_required:
        review_decisions, latest_review_decision = append_review_decision(
            document,
            run_id=run_doc["id"],
            decision_type="flag",
            rationale="Classification confidence fell below the configured auto-accept threshold.",
            resulting_state="awaiting_review",
        )
        current_review_state = "awaiting_review"

    is_academic = suggested_category == "Academic Papers" or "My Writings" in suggested_category
    updated_document = {
        "ai_suggested_category": suggested_category,
        "category": suggested_category,
        "document_type": classification.get("document_type"),
        "classification_confidence": classification.get("confidence"),
        "classification_rationale": classification.get("rationale"),
        "is_academic": is_academic,
        "review_required": review_required,
        "current_state": "classified",
        "current_review_state": current_review_state,
        "processing_runs": processing_runs,
        "review_decision": latest_review_decision,
        "review_decisions": review_decisions,
        "latest_run_id": run_doc["id"],
        "latest_job_id": job["id"],
        "audit_log": append_audit_event(
            document,
            "categorize",
            {
                "job_id": job["id"],
                "category": suggested_category,
                "document_type": classification.get("document_type"),
                "confidence": classification.get("confidence"),
            },
        ),
    }
    await db.documents.update_one({"id": document["id"]}, {"$set": updated_document})
    await persist_processing_run_record(document["id"], run_doc)
    return {
        "document_id": document["id"],
        "suggested_category": suggested_category,
        "document_type": classification.get("document_type"),
        "confidence": classification.get("confidence"),
    }


async def process_summary_job(document: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    text_result = get_document_text_for_processing(document, max_pages=50)
    text = text_result.get("text", "")
    if not text or text_result.get("text_quality") != "usable":
        raise no_text_http_exception(text_result, action="summarize")

    summary = await ai_generate_summary(text)
    processing_runs, run_doc = append_processing_run(
        document,
        ProcessingJobType.SUMMARY.value,
        triggered_by="job_worker",
        status="completed",
        job_id=job["id"],
        started_at=job.get("started_at"),
        completed_at=iso_now(),
        details={"summary_length": len(summary)},
    )
    await db.documents.update_one(
        {"id": document["id"]},
        {"$set": {
            "summary": summary,
            "current_state": "summarized",
            "processing_runs": processing_runs,
            "latest_run_id": run_doc["id"],
            "latest_job_id": job["id"],
            "audit_log": append_audit_event(document, "summary", {"job_id": job["id"], "summary_length": len(summary)}),
        }},
    )
    await persist_processing_run_record(document["id"], run_doc)
    return {"document_id": document["id"], "summary": summary}


async def process_extract_job(document: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    text_result = get_document_text_for_processing(document, max_pages=50)
    text = text_result.get("text", "")
    if not text or text_result.get("text_quality") != "usable":
        raise no_text_http_exception(text_result, action="extract structured fields from")

    extraction = await ai_extract_fields(text)
    fields = extraction.get("fields", {}) if isinstance(extraction.get("fields"), dict) else {}
    controls = run_control_checks(fields)
    review_required = bool(controls["missing_fields"] or controls["invalid_values"])
    processing_runs, run_doc = append_processing_run(
        document,
        ProcessingJobType.EXTRACT.value,
        triggered_by="job_worker",
        status="completed",
        job_id=job["id"],
        started_at=job.get("started_at"),
        completed_at=iso_now(),
        details={
            "field_count": len(fields),
            "missing_count": len(controls["missing_fields"]),
            "invalid_count": len(controls["invalid_values"]),
        },
    )
    review_decisions = list(document.get("review_decisions", []))
    latest_review_decision = document.get("review_decision")
    current_review_state = "extraction_complete"
    if review_required:
        review_decisions, latest_review_decision = append_review_decision(
            document,
            run_id=run_doc["id"],
            decision_type="flag",
            rationale="Control checks identified missing or invalid fields.",
            resulting_state="awaiting_review",
        )
        current_review_state = "awaiting_review"
    else:
        review_decisions, latest_review_decision = append_review_decision(
            document,
            run_id=run_doc["id"],
            decision_type="approve",
            rationale="Extraction passed configured control checks without requiring reviewer intervention.",
            resulting_state="reviewed_approved",
        )
        current_review_state = "reviewed_approved"

    compliance_note = (
        "Contains potentially sensitive data; apply restricted access and masking where required."
        if any(tag in text.lower() for tag in ["patient", "ssn", "iban", "account", "medical"]) else
        "No obvious sensitive marker found; enforce standard retention and least-privilege access."
    )
    artifact_versions, _ = append_artifact_version(
        document,
        source="extracted",
        data={
            "document_type": extraction.get("document_type") or document.get("document_type"),
            "fields": fields,
            "control_checks": controls,
        },
    )
    updated_document = {
        "document_type": extraction.get("document_type") or document.get("document_type"),
        "extraction": extraction,
        "control_checks": controls,
        "review_required": review_required,
        "compliance_note": compliance_note,
        "current_state": "extracted",
        "current_review_state": current_review_state,
        "processing_runs": processing_runs,
        "review_decision": latest_review_decision,
        "review_decisions": review_decisions,
        "artifact_versions": artifact_versions,
        "latest_run_id": run_doc["id"],
        "latest_job_id": job["id"],
        "audit_log": append_audit_event(
            document,
            "extract",
            {
                "job_id": job["id"],
                "field_count": len(fields),
                "missing_count": len(controls["missing_fields"]),
                "invalid_count": len(controls["invalid_values"]),
            },
        ),
    }
    await db.documents.update_one({"id": document["id"]}, {"$set": updated_document})
    await persist_processing_run_record(document["id"], run_doc)
    return {
        "document_id": document["id"],
        "document_type": updated_document["document_type"],
        "data": extraction,
        "review_required": review_required,
        "controls": controls,
    }


async def process_citations_job(document: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    text_result = get_document_text_for_processing(document, max_pages=50)
    text = text_result.get("text", "")
    if not text or text_result.get("text_quality") != "usable":
        raise no_text_http_exception(text_result, action="extract citations from")

    citations = await ai_extract_citations(text)
    processing_runs, run_doc = append_processing_run(
        document,
        ProcessingJobType.CITATIONS.value,
        triggered_by="job_worker",
        status="completed",
        job_id=job["id"],
        started_at=job.get("started_at"),
        completed_at=iso_now(),
        details={"citation_count": len(citations)},
    )
    await db.documents.update_one(
        {"id": document["id"]},
        {"$set": {
            "citations": citations,
            "current_state": "citations_extracted",
            "processing_runs": processing_runs,
            "latest_run_id": run_doc["id"],
            "latest_job_id": job["id"],
            "audit_log": append_audit_event(document, "citations", {"job_id": job["id"], "citation_count": len(citations)}),
        }},
    )
    await persist_processing_run_record(document["id"], run_doc)
    return {"document_id": document["id"], "citations": citations}


async def execute_processing_job(job: Dict[str, Any]) -> Dict[str, Any]:
    document = await db.documents.find_one({"id": job["document_id"]}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    job_type = job.get("job_type")
    if job_type == ProcessingJobType.OCR.value:
        return await process_ocr_job(document, job)
    if job_type == ProcessingJobType.CLASSIFY.value:
        return await process_classify_job(document, job)
    if job_type == ProcessingJobType.SUMMARY.value:
        return await process_summary_job(document, job)
    if job_type == ProcessingJobType.EXTRACT.value:
        return await process_extract_job(document, job)
    if job_type == ProcessingJobType.CITATIONS.value:
        return await process_citations_job(document, job)
    raise HTTPException(status_code=400, detail=f"Unsupported job type: {job_type}")


async def run_processing_job(job: Dict[str, Any]) -> None:
    try:
        output = await execute_processing_job(job)
        await mark_job_result(job, status=ProcessingJobStatus.SUCCESS, output=output)
        requested_by = None
        if job.get("requested_by_user_id"):
            requested_by = {
                "id": job.get("requested_by_user_id"),
                "role": job.get("requested_by_role"),
            }
        for next_job_type in job.get("next_job_types", []):
            try:
                await enqueue_processing_job(job["document_id"], ProcessingJobType(next_job_type), requested_by)
            except Exception as exc:  # pragma: no cover - best-effort chaining
                await append_document_audit(
                    job["document_id"],
                    "job_chain_failed",
                    {"parent_job_id": job["id"], "next_job_type": next_job_type, "error": str(exc)},
                )
    except HTTPException as exc:
        await mark_job_result(job, status=ProcessingJobStatus.FAILED, error_message=exc.detail)
    except Exception as exc:  # pragma: no cover - defensive worker protection
        await mark_job_result(job, status=ProcessingJobStatus.FAILED, error_message=str(exc))


async def processing_job_worker(stop_event: asyncio.Event) -> None:
    while not stop_event.is_set():
        job = await claim_next_processing_job()
        if not job:
            try:
                await asyncio.wait_for(stop_event.wait(), timeout=AURORAI_JOB_POLL_INTERVAL_SECONDS)
            except asyncio.TimeoutError:
                continue
            continue
        await run_processing_job(job)


async def create_document_record(
    original_filename: str,
    content: bytes,
) -> Dict[str, Any]:
    allowed_extensions = SUPPORTED_UPLOAD_EXTENSIONS
    file_ext = Path(original_filename).suffix.lower()

    if file_ext == ".doc":
        raise HTTPException(status_code=400, detail="Legacy DOC upload is not supported. Use PDF, TXT, or DOCX.")
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"File type {file_ext} not supported. Allowed: {allowed_extensions}")

    doc_id = str(uuid.uuid4())
    new_filename = f"{doc_id}{file_ext}"
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    file_path = UPLOAD_DIR / new_filename

    async with aiofiles.open(file_path, "wb") as handle:
        await handle.write(content)

    try:
        extraction_result = extract_document_content(file_path, raw_bytes=content, allow_ocr=False)
    except DocumentExtractionError as exc:
        raise HTTPException(status_code=exc.status_code, detail=exc.message) from exc

    text_preview = extraction_result.get("text", "")[:5000] if extraction_result.get("text") else ""
    ingestion_details = {
        "text_source": extraction_result.get("text_source"),
        "text_quality": extraction_result.get("text_quality"),
        "ocr_status": extraction_result.get("ocr_status"),
        "ocr_reason": extraction_result.get("ocr_reason"),
        "warnings": extraction_result.get("warnings", []),
    }
    current_state = (
        "uploaded_queued_for_ocr"
        if ingestion_details["ocr_status"] == "queued"
        else "uploaded_requires_ocr"
        if ingestion_details["ocr_status"] == "required_unavailable"
        else "uploaded"
    )
    current_review_state = (
        "awaiting_ocr"
        if ingestion_details["ocr_status"] == "queued"
        else "ocr_unavailable"
        if ingestion_details["ocr_status"] == "required_unavailable"
        else "pending"
    )
    review_required = ingestion_details["ocr_status"] == "required_unavailable"
    base_document: Dict[str, Any] = {}
    artifact_versions, _ = append_artifact_version(
        base_document,
        source="original",
        data={
            "filename": original_filename,
            "file_size": len(content),
            "source_hash": compute_sha256_bytes(content),
            "text_preview": text_preview[:2000] if text_preview else "",
            "ingestion_details": ingestion_details,
        },
    )
    document = Document(
        id=doc_id,
        filename=new_filename,
        original_filename=original_filename,
        file_size=len(content),
        page_count=extraction_result.get("page_count", 0),
        text_preview=text_preview[:2000] if text_preview else "",
        source_hash=compute_sha256_bytes(content),
        ingestion_details=ingestion_details,
        current_state=current_state,
        current_review_state=current_review_state,
        review_required=review_required,
        artifact_versions=artifact_versions,
    )
    doc_dict = document.model_dump()
    doc_dict["uploaded_at"] = doc_dict["uploaded_at"].isoformat()
    doc_dict["audit_log"] = append_audit_event(
        doc_dict,
        "upload",
        {
            "original_filename": original_filename,
            "file_size": len(content),
            "page_count": extraction_result.get("page_count", 0),
            "source_hash": doc_dict["source_hash"],
            "text_source": ingestion_details["text_source"],
            "text_quality": ingestion_details["text_quality"],
            "ocr_status": ingestion_details["ocr_status"],
            "ocr_reason": ingestion_details["ocr_reason"],
        },
    )
    await db.documents.insert_one(doc_dict)
    return doc_dict


# Routes
class BulkCategorizeRequest(BaseModel):
    document_ids: List[str]


def summarize_job_counts(jobs: List[Dict[str, Any]]) -> Dict[str, int]:
    counts = {
        ProcessingJobStatus.QUEUED.value: 0,
        ProcessingJobStatus.RUNNING.value: 0,
        ProcessingJobStatus.SUCCESS.value: 0,
        ProcessingJobStatus.FAILED.value: 0,
    }
    for job in jobs:
        status = job.get("status")
        if status in counts:
            counts[status] += 1
    return counts


def build_review_flags(document: Dict[str, Any]) -> List[Dict[str, Any]]:
    flags: List[Dict[str, Any]] = []
    if document.get("review_required"):
        flags.append({"type": "review_required", "label": "Human review required"})

    confidence = document.get("classification_confidence")
    if isinstance(confidence, (int, float)) and float(confidence) < 0.75:
        flags.append({
            "type": "classification_confidence",
            "label": "Low classification confidence",
            "value": round(float(confidence), 4),
        })

    ingestion_details = document.get("ingestion_details", {}) or {}
    if ingestion_details.get("ocr_status") in {"required_unavailable", "failed"}:
        flags.append({
            "type": "ocr",
            "label": "OCR requires attention",
            "value": ingestion_details.get("ocr_reason"),
        })

    for key in ("missing_fields", "invalid_values", "duplicate_values", "inconsistencies"):
        values = document.get("control_checks", {}).get(key, [])
        if values:
            flags.append({
                "type": key,
                "label": key.replace("_", " "),
                "count": len(values),
            })

    return flags


async def list_document_jobs(document_id: str, *, limit: int = 100) -> List[Dict[str, Any]]:
    return await db.processing_jobs.find({"document_id": document_id}, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)


async def build_document_status_payload(document: Dict[str, Any]) -> Dict[str, Any]:
    jobs = await list_document_jobs(document["id"], limit=100)
    job_counts = summarize_job_counts(jobs)
    latest_job = jobs[0] if jobs else None
    return {
        "document": document,
        "latest_job": latest_job,
        "jobs": jobs,
        "job_counts": job_counts,
        "has_active_job": job_counts[ProcessingJobStatus.QUEUED.value] > 0 or job_counts[ProcessingJobStatus.RUNNING.value] > 0,
        "review_flags": build_review_flags(document),
    }


async def maybe_enqueue_initial_ocr_job(document: Dict[str, Any], user: Dict[str, Any]) -> List[Dict[str, Any]]:
    if (document.get("ingestion_details", {}) or {}).get("ocr_status") != "queued":
        return []
    return [await enqueue_processing_job(document["id"], ProcessingJobType.OCR, user)]


async def ensure_document_exists(doc_id: str) -> Dict[str, Any]:
    document = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document


async def queue_document_job(
    document_id: str,
    job_type: ProcessingJobType,
    user: Dict[str, Any],
    *,
    next_job_types: Optional[List[ProcessingJobType]] = None,
) -> Dict[str, Any]:
    document = await ensure_document_exists(document_id)
    ingestion_details = document.get("ingestion_details", {}) or {}
    capabilities = build_runtime_capabilities()

    if job_type == ProcessingJobType.OCR and not capabilities["ocr_available"]:
        raise HTTPException(status_code=503, detail=f"OCR runtime is unavailable ({capabilities['ocr_reason']}).")

    if job_type != ProcessingJobType.OCR and ingestion_details.get("ocr_status") == "required_unavailable":
        raise HTTPException(
            status_code=503,
            detail=f"OCR is required before AurorA can {job_type.value} this document, but the runtime is unavailable ({ingestion_details.get('ocr_reason')}).",
        )

    if job_type != ProcessingJobType.OCR and ingestion_details.get("ocr_status") == "queued":
        existing_jobs = await list_document_jobs(document_id, limit=50)
        has_existing_ocr = any(job.get("job_type") == ProcessingJobType.OCR.value for job in existing_jobs)
        if not has_existing_ocr:
            job = await enqueue_processing_job(document_id, ProcessingJobType.OCR, user, next_job_types=[job_type])
            return {
                "accepted": True,
                "queued_via": "ocr_preflight",
                "document_id": document_id,
                "job": job,
            }

    job = await enqueue_processing_job(document_id, job_type, user, next_job_types=next_job_types)
    return {
        "accepted": True,
        "document_id": document_id,
        "job": job,
    }


@api_router.get("/")
async def root():
    return {
        "message": "AurorAI IDP API",
        "mission": IDP_MISSION,
        "pipeline": IDP_PIPELINE,
        "capabilities": build_runtime_capabilities(),
    }


@api_router.get("/idp/pipeline")
async def get_idp_pipeline():
    return {
        "mission": IDP_MISSION,
        "pipeline": IDP_PIPELINE,
        "capabilities": build_runtime_capabilities(),
        "golden_rules": [
            "Always return one human-readable output and one machine-readable output.",
            "Always include confidence, ambiguity flags, and recommended next actions.",
            "Focus on risk-aware processing for regulated sectors.",
        ],
    }


@api_router.get("/categories")
async def get_categories():
    return {"categories": CATEGORIES}


@api_router.get("/stats")
async def get_stats():
    pipeline = [
        {"$group": {"_id": "$category", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ]
    results = await db.documents.aggregate(pipeline).to_list(100)
    total = sum(result["count"] for result in results)

    stats = [
        {
            "category": result["_id"],
            "count": result["count"],
            "percentage": round((result["count"] / total * 100) if total > 0 else 0, 1),
        }
        for result in results
    ]

    existing_categories = {item["category"] for item in stats}
    for category in CATEGORIES:
        if category not in existing_categories:
            stats.append({"category": category, "count": 0, "percentage": 0})

    return {"stats": stats, "total_documents": total}


@api_router.get("/config")
@api_router.get("/internal/config")
async def get_config(request: Request, current_user: Optional[Dict[str, Any]] = Depends(get_current_user)):
    origin = str(request.base_url).rstrip("/")
    return {
        "surface": "aurorai",
        "origin": origin,
        "proxy_path": "/internal-api/aurorai",
        "api_base_path": "/internal-api/aurorai/api",
        "job_poll_interval_seconds": AURORAI_JOB_POLL_INTERVAL_SECONDS,
        "capabilities": build_runtime_capabilities(),
        "session": {
            "authenticated": bool(current_user),
            "user": serialize_user(current_user),
        },
        "available_roles": [role.value for role in UserRole],
    }


@api_router.post("/session/login")
@api_router.post("/internal/session/login")
async def login_session(payload: UserLogin, response: Response):
    email = normalize_email(payload.email)
    user = await db.users.find_one({"email": email, "is_active": True}, {"_id": 0})
    if not user or not verify_password(payload.password, user.get("hashed_password", "")):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    token = create_session_token({"sub": user["id"], "role": user.get("role")})
    set_session_cookie(response, token)
    return SessionToken(access_token=token, user=serialize_user(user)).model_dump()


@api_router.post("/session/logout")
@api_router.post("/internal/session/logout")
async def logout_session(response: Response):
    clear_session_cookie(response)
    return {"signed_out": True}


@api_router.get("/session/me")
@api_router.get("/internal/session/me")
async def get_session_me(current_user: Dict[str, Any] = Depends(require_auth_user)):
    return {"authenticated": True, "user": serialize_user(current_user)}


@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    content = await file.read()
    document = await create_document_record(file.filename or "upload", content)
    queued_jobs = await maybe_enqueue_initial_ocr_job(document, current_user)
    log_pipeline_event(
        "document_uploaded",
        document_id=document["id"],
        filename=file.filename,
        file_ext=Path(file.filename or "").suffix.lower(),
        text_source=(document.get("ingestion_details", {}) or {}).get("text_source"),
        ocr_status=(document.get("ingestion_details", {}) or {}).get("ocr_status"),
    )
    return {
        **document,
        "queued_jobs": queued_jobs,
    }


@api_router.post("/upload/batch")
async def upload_batch(
    archive: Optional[UploadFile] = File(None),
    files: Optional[List[UploadFile]] = File(None),
    folder_reference: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    if archive is None and not files:
        raise HTTPException(status_code=400, detail="Provide a zip archive or one or more files.")

    uploads: List[tuple[str, bytes]] = []
    if archive is not None:
        if Path(archive.filename or "").suffix.lower() != ".zip":
            raise HTTPException(status_code=400, detail="Batch archive uploads must be provided as a .zip file.")
        try:
            payload = await archive.read()
            with ZipFile(BytesIO(payload)) as zip_file:
                for item in zip_file.infolist():
                    if item.is_dir():
                        continue
                    uploads.append((Path(item.filename).name, zip_file.read(item.filename)))
        except BadZipFile as exc:
            raise HTTPException(status_code=400, detail=f"Could not read zip archive: {exc}") from exc

    for upload in files or []:
        uploads.append((upload.filename or "upload", await upload.read()))

    created_documents = []
    queued_jobs = []
    errors = []
    for filename, content in uploads:
        try:
            document = await create_document_record(filename, content)
            created_documents.append(document)
            queued_jobs.extend(await maybe_enqueue_initial_ocr_job(document, current_user))
        except HTTPException as exc:
            errors.append({"filename": filename, "error": exc.detail, "status_code": exc.status_code})

    return {
        "folder_reference": folder_reference,
        "created_count": len(created_documents),
        "failed_count": len(errors),
        "documents": created_documents,
        "queued_jobs": queued_jobs,
        "errors": errors,
    }


@api_router.post("/process", status_code=202)
async def create_process_job(
    request: ProcessingJobRequest,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    return await queue_document_job(
        request.document_id,
        request.job_type,
        current_user,
        next_job_types=request.next_job_types or None,
    )


@api_router.get("/process/{job_id}")
async def get_process_job(job_id: str, _current_user: Dict[str, Any] = Depends(require_document_read_access)):
    job = await db.processing_jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Processing job not found")

    document = await db.documents.find_one({"id": job["document_id"]}, {"_id": 0})
    return {
        "job": job,
        "document": document,
    }


@api_router.get("/status/overview")
async def get_status_overview(_current_user: Dict[str, Any] = Depends(require_document_read_access)):
    documents = await db.documents.find({}, {"_id": 0}).to_list(5000)
    jobs = await db.processing_jobs.find({}, {"_id": 0}).to_list(5000)
    job_counts = summarize_job_counts(jobs)
    review_pending = [
        document for document in documents
        if document.get("review_required") or document.get("current_review_state") in {"awaiting_review", "ocr_unavailable", "review_flagged"}
    ]

    return {
        "documents": {
            "total": len(documents),
            "review_pending": len(review_pending),
        },
        "jobs": job_counts,
        "recent_failed_jobs": [job for job in jobs if job.get("status") == ProcessingJobStatus.FAILED.value][:10],
    }


@api_router.get("/status/{doc_id}")
async def get_document_status(doc_id: str, _current_user: Dict[str, Any] = Depends(require_document_read_access)):
    document = await ensure_document_exists(doc_id)
    return await build_document_status_payload(document)


@api_router.get("/queue/review")
async def get_review_queue(
    status: str = Query("pending"),
    assigned_to: Optional[str] = Query(None),
    current_user: Dict[str, Any] = Depends(require_reviewer_access),
):
    documents = await db.documents.find({}, {"_id": 0}).sort("uploaded_at", -1).to_list(1000)

    def matches_status(document: Dict[str, Any]) -> bool:
        review_state = document.get("current_review_state")
        pending = document.get("review_required") or review_state in {"awaiting_review", "ocr_unavailable", "review_flagged"}
        reviewed = review_state in {"reviewed_approved", "reviewed_rejected"}
        if status == "pending":
            return pending
        if status == "reviewed":
            return reviewed
        return True

    queue = []
    for document in documents:
        if not matches_status(document):
            continue
        if assigned_to == "me":
            assigned_user_id = document.get("assigned_reviewer_user_id")
            if assigned_user_id not in {None, current_user["id"]}:
                continue

        queue.append({
            "document": document,
            "flags": build_review_flags(document),
            "thumbnail": {
                "label": document.get("original_filename"),
                "file_extension": Path(document.get("original_filename", "")).suffix.lower(),
            },
        })

    return {"documents": queue, "count": len(queue)}


@api_router.post("/review/{doc_id}")
async def submit_review_decision(
    doc_id: str,
    payload: ReviewDecisionRequest,
    current_user: Dict[str, Any] = Depends(require_reviewer_access),
):
    document = await ensure_document_exists(doc_id)
    state_map = {
        "approve": ("reviewed", "reviewed_approved"),
        "reject": ("reviewed", "reviewed_rejected"),
        "flag": ("review_flagged", "awaiting_review"),
    }
    current_state, current_review_state = state_map[payload.decision]
    review_required = payload.decision == "flag"
    review_decisions, review_decision = append_review_decision(
        document,
        run_id=document.get("latest_run_id"),
        decision_type=payload.decision,
        rationale=payload.comment or f"Reviewer marked the document as {payload.decision}.",
        resulting_state=current_review_state,
        actor_type="user",
        actor_id=current_user["id"],
        reviewer_user_id=current_user["id"],
        comment=payload.comment,
    )
    updated_audit = append_audit_event(
        document,
        "review",
        {
            "reviewer_user_id": current_user["id"],
            "decision": payload.decision,
            "comment": payload.comment,
        },
    )

    await db.documents.update_one(
        {"id": doc_id},
        {"$set": {
            "review_required": review_required,
            "current_state": current_state,
            "current_review_state": current_review_state,
            "assigned_reviewer_user_id": current_user["id"],
            "review_decision": review_decision,
            "review_decisions": review_decisions,
            "audit_log": updated_audit,
        }},
    )

    updated_document = await ensure_document_exists(doc_id)
    return {
        "document": updated_document,
        "review_decision": review_decision,
    }


@api_router.post("/reprocess/{doc_id}", status_code=202)
async def reprocess_document(
    doc_id: str,
    job_type: ProcessingJobType = Query(...),
    current_user: Dict[str, Any] = Depends(require_reviewer_access),
):
    return await queue_document_job(doc_id, job_type, current_user)


@api_router.post("/documents/{doc_id}/categorize", status_code=202)
async def categorize_document(
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    return await queue_document_job(doc_id, ProcessingJobType.CLASSIFY, current_user)


@api_router.post("/documents/bulk-categorize", status_code=202)
async def bulk_categorize_documents(
    request: BulkCategorizeRequest,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    results = []
    for document_id in request.document_ids:
        try:
            results.append(await queue_document_job(document_id, ProcessingJobType.CLASSIFY, current_user))
        except HTTPException as exc:
            results.append({"document_id": document_id, "error": exc.detail, "status_code": exc.status_code})
    return {"results": results}


@api_router.post("/documents/{doc_id}/summary", status_code=202)
async def generate_summary(
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    return await queue_document_job(doc_id, ProcessingJobType.SUMMARY, current_user)


@api_router.post("/documents/{doc_id}/extract", status_code=202)
async def extract_fields(
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    return await queue_document_job(doc_id, ProcessingJobType.EXTRACT, current_user)


@api_router.post("/documents/{doc_id}/citations", status_code=202)
async def extract_citations(
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    return await queue_document_job(doc_id, ProcessingJobType.CITATIONS, current_user)


@api_router.post("/documents/{doc_id}/evidence-package")
async def generate_evidence_package(
    doc_id: str,
    request: EvidencePackageRequest,
    _current_user: Dict[str, Any] = Depends(require_operator_access),
):
    evidence_record = await create_evidence_record(doc_id, request)
    evidence_record.pop("document_audit_log", None)
    return evidence_record


@api_router.post("/documents/{doc_id}/handoff-to-compassai")
async def handoff_to_compassai(
    doc_id: str,
    request: EvidenceHandoffRequest,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    if not COMPASSAI_INGEST_TOKEN:
        raise HTTPException(status_code=503, detail="COMPASSAI_INGEST_TOKEN is not configured on the server.")

    evidence_record = await create_evidence_record(doc_id, request)
    evidence_record.pop("document_audit_log", None)
    target_base_url = (request.compassai_base_url or COMPASSAI_BASE_URL).rstrip("/")
    response = await post_json(
        f"{target_base_url}/api/v1/evidence",
        {
            "usecase_id": evidence_record["usecase_id"],
            "producer": evidence_record["producer"],
            "artifact_type": evidence_record["artifact_type"],
            "payload": evidence_record["payload"],
            "hash": evidence_record["hash"],
            "control_checks": evidence_record["control_checks"],
        },
        headers={"Authorization": f"Bearer {COMPASSAI_INGEST_TOKEN}"},
    )

    document = await ensure_document_exists(doc_id)
    handoff_history, handoff_doc = append_handoff_history(
        document,
        evidence_id=evidence_record["id"],
        target="compassai",
        target_url=f"{target_base_url}/api/v1/evidence",
        status_code=response["status_code"],
    )
    updated_audit = append_audit_event(
        document,
        "handoff_to_compassai",
        {
            "user_id": current_user["id"],
            "target_url": f"{target_base_url}/api/v1/evidence",
            "status_code": response["status_code"],
            "evidence_id": evidence_record["id"],
        },
    )

    await db.evidence_handoffs.insert_one(
        {
            "id": handoff_doc["id"],
            "document_id": doc_id,
            "evidence_id": evidence_record["id"],
            "usecase_id": request.usecase_id,
            "target": "compassai",
            "target_url": f"{target_base_url}/api/v1/evidence",
            "status_code": response["status_code"],
            "response": response["body"],
            "created_at": iso_now(),
        }
    )

    await db.documents.update_one(
        {"id": doc_id},
        {"$set": {
            "handoff_history": handoff_history,
            "latest_handoff_id": handoff_doc["id"],
            "current_state": "handed_off_to_compassai" if response["status_code"] < 400 else "handoff_failed",
            "audit_log": updated_audit,
        }},
    )

    if response["status_code"] >= 400:
        raise HTTPException(
            status_code=response["status_code"],
            detail={
                "message": "CompassAI rejected the evidence handoff.",
                "compassai_response": response["body"],
                "evidence_id": evidence_record["id"],
            },
        )

    return {
        "message": "Evidence handed off to CompassAI.",
        "evidence": evidence_record,
        "compassai_response": response["body"],
    }


@api_router.get("/documents")
async def get_documents(
    category: Optional[str] = None,
    search: Optional[str] = None,
    is_academic: Optional[bool] = None,
    reading_list_id: Optional[str] = None,
    _current_user: Dict[str, Any] = Depends(require_document_read_access),
):
    query: Dict[str, Any] = {}
    if category:
        query["category"] = category
    if is_academic is not None:
        query["is_academic"] = is_academic
    if reading_list_id:
        query["reading_list_id"] = reading_list_id
    if search:
        query["$or"] = [
            {"original_filename": {"$regex": search, "$options": "i"}},
            {"text_preview": {"$regex": search, "$options": "i"}},
        ]

    documents = await db.documents.find(query, {"_id": 0}).sort("uploaded_at", -1).to_list(500)
    return {"documents": documents}


@api_router.get("/documents/{doc_id}")
async def get_document(
    doc_id: str,
    _current_user: Dict[str, Any] = Depends(require_document_read_access),
):
    return await ensure_document_exists(doc_id)


@api_router.get("/documents/{doc_id}/download")
async def download_document(
    doc_id: str,
    _current_user: Dict[str, Any] = Depends(require_document_read_access),
):
    document = await ensure_document_exists(doc_id)
    file_path = UPLOAD_DIR / document["filename"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on server")

    return FileResponse(
        path=str(file_path),
        filename=document["original_filename"],
        media_type="application/octet-stream",
    )


@api_router.patch("/documents/{doc_id}")
async def update_document(
    doc_id: str,
    update: DocumentUpdate,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    document = await ensure_document_exists(doc_id)
    update_data = {key: value for key, value in update.model_dump().items() if value is not None}
    if "category" in update_data:
        update_data["is_academic"] = update_data["category"] in {"Academic Papers", "My Writings & Publications"}

    if update_data:
        update_data["audit_log"] = append_audit_event(
            document,
            "document_update",
            {
                "user_id": current_user["id"],
                "fields": sorted(update_data.keys()),
            },
        )
        await db.documents.update_one({"id": doc_id}, {"$set": update_data})

    return await ensure_document_exists(doc_id)


@api_router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    document = await ensure_document_exists(doc_id)
    file_path = UPLOAD_DIR / document["filename"]
    if file_path.exists():
        file_path.unlink()

    await db.deleted_documents.insert_one(
        {
            "document_id": doc_id,
            "deleted_at": iso_now(),
            "deleted_by_user_id": current_user["id"],
            "document": document,
        }
    )
    await db.documents.delete_one({"id": doc_id})
    return {"message": "Document deleted", "id": doc_id}


@api_router.post("/reading-lists")
async def create_reading_list(
    data: ReadingListCreate,
    _current_user: Dict[str, Any] = Depends(require_operator_access),
):
    reading_list = ReadingList(name=data.name, description=data.description)
    reading_list_doc = reading_list.model_dump()
    reading_list_doc["created_at"] = reading_list_doc["created_at"].isoformat()
    await db.reading_lists.insert_one(reading_list_doc)
    return reading_list_doc


@api_router.get("/reading-lists")
async def get_reading_lists(_current_user: Dict[str, Any] = Depends(require_document_read_access)):
    reading_lists = await db.reading_lists.find({}, {"_id": 0}).to_list(100)
    for reading_list in reading_lists:
        reading_list["document_count"] = await db.documents.count_documents({"reading_list_id": reading_list["id"]})
    return {"reading_lists": reading_lists}


@api_router.get("/reading-lists/{list_id}")
async def get_reading_list(
    list_id: str,
    _current_user: Dict[str, Any] = Depends(require_document_read_access),
):
    reading_list = await db.reading_lists.find_one({"id": list_id}, {"_id": 0})
    if not reading_list:
        raise HTTPException(status_code=404, detail="Reading list not found")

    reading_list["documents"] = await db.documents.find({"reading_list_id": list_id}, {"_id": 0}).to_list(100)
    return reading_list


@api_router.post("/reading-lists/{list_id}/documents/{doc_id}")
async def add_to_reading_list(
    list_id: str,
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    reading_list = await db.reading_lists.find_one({"id": list_id}, {"_id": 0})
    if not reading_list:
        raise HTTPException(status_code=404, detail="Reading list not found")

    document = await ensure_document_exists(doc_id)
    await db.documents.update_one(
        {"id": doc_id},
        {"$set": {
            "reading_list_id": list_id,
            "audit_log": append_audit_event(
                document,
                "reading_list_add",
                {"list_id": list_id, "user_id": current_user["id"]},
            ),
        }},
    )
    return {"message": "Document added to reading list"}


@api_router.delete("/reading-lists/{list_id}/documents/{doc_id}")
async def remove_from_reading_list(
    list_id: str,
    doc_id: str,
    current_user: Dict[str, Any] = Depends(require_operator_access),
):
    document = await ensure_document_exists(doc_id)
    await db.documents.update_one(
        {"id": doc_id, "reading_list_id": list_id},
        {"$set": {
            "reading_list_id": None,
            "audit_log": append_audit_event(
                document,
                "reading_list_remove",
                {"list_id": list_id, "user_id": current_user["id"]},
            ),
        }},
    )
    return {"message": "Document removed from reading list"}


@api_router.delete("/reading-lists/{list_id}")
async def delete_reading_list(
    list_id: str,
    _current_user: Dict[str, Any] = Depends(require_operator_access),
):
    reading_list = await db.reading_lists.find_one({"id": list_id}, {"_id": 0})
    if not reading_list:
        raise HTTPException(status_code=404, detail="Reading list not found")

    await db.documents.update_many({"reading_list_id": list_id}, {"$set": {"reading_list_id": None}})
    await db.reading_lists.delete_one({"id": list_id})
    return {"message": "Reading list deleted"}


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


@asynccontextmanager
async def aurorai_lifespan(_app: FastAPI):
    stop_event = asyncio.Event()
    worker_task: Optional[asyncio.Task[Any]] = None
    try:
        await seed_bootstrap_users()
        await backfill_document_schema()
        enforce_required_runtime()
        worker_task = asyncio.create_task(processing_job_worker(stop_event))
        _app.state.processing_stop_event = stop_event
        _app.state.processing_worker_task = worker_task
        yield
    finally:
        stop_event.set()
        if worker_task is not None:
            worker_task.cancel()
            try:
                await worker_task
            except asyncio.CancelledError:
                pass
        client.close()


app.router.lifespan_context = aurorai_lifespan
